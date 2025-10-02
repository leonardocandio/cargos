"""
Service classes for Excel processing and file generation.
"""
import pandas as pd
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from models import ExcelData, GenerationResult, AppConfig, ExcelValidationResult, WorksheetMetadata, WorksheetParsingResult, GenerationOptions
from validators import TemplateValidator
from unified_config_service import UnifiedConfigService
from constants import (
    METADATA_ROW_FECHA_SOLICITUD, METADATA_COL_FECHA_SOLICITUD,
    METADATA_ROW_TIENDA, METADATA_COL_TIENDA,
    METADATA_ROW_ADMINISTRADOR, METADATA_COL_ADMINISTRADOR,
    DATA_START_ROW, IGNORE_COLUMN_INDEX, MAIN_DATA_END_COLUMN,
    UNIFORM_DATA_START_ROW, UNIFORM_DATA_START_COLUMN, UNIFORM_DATA_END_COLUMN,
    UNIFORM_COLUMN_NAMES, SPANISH_MONTHS
)


class ExcelService:
    """Service for handling Excel file operations."""
    
    def __init__(self, logger: logging.Logger):
        self.logger = logger
    
    def load_excel_file(self, file_path: str) -> ExcelData:
        """
        Load Excel file and parse all worksheets.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            ExcelData object with parsed worksheets or empty if failed
        """
        try:
            if not file_path:
                raise ValueError("File path is empty")
            
            if not Path(file_path).exists():
                raise FileNotFoundError(f"File does not exist: {file_path}")
            
            self.logger.info(f"Loading Excel file: {file_path}")
            
            # Read all sheets from Excel file
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            self.logger.info(f"Found {len(sheet_names)} worksheets: {sheet_names}")
            
            worksheets = []
            for sheet_name in sheet_names:
                worksheet_result = self._parse_worksheet(excel_file, sheet_name)
                worksheets.append(worksheet_result)
            
            excel_data = ExcelData(
                file_path=file_path,
                worksheets=worksheets
            )
            
            self.logger.info(f"Excel file loaded successfully. {excel_data.successful_worksheets}/{excel_data.total_worksheets} worksheets parsed, {excel_data.total_people_parsed} people total.")
            
            return excel_data
            
        except Exception as e:
            error_msg = f"Error loading Excel file: {str(e)}"
            self.logger.exception(error_msg)
            raise Exception(error_msg)
    
    def _parse_worksheet(self, excel_file: pd.ExcelFile, sheet_name: str) -> WorksheetParsingResult:
        """
        Parse a single worksheet and extract metadata and data.
        
        Args:
            excel_file: Pandas ExcelFile object
            sheet_name: Name of the sheet to parse
            
        Returns:
            WorksheetParsingResult with parsed data and metadata
        """
        metadata = WorksheetMetadata(sheet_name=sheet_name)
        result = WorksheetParsingResult(metadata=metadata)
        
        try:
            # Read the entire sheet first to extract metadata
            sheet_data = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
            
            # Extract metadata from specific cells
            try:
                if len(sheet_data) > METADATA_ROW_FECHA_SOLICITUD and len(sheet_data.columns) > METADATA_COL_FECHA_SOLICITUD:
                    metadata.fecha_solicitud = str(sheet_data.iloc[METADATA_ROW_FECHA_SOLICITUD, METADATA_COL_FECHA_SOLICITUD]) if pd.notna(sheet_data.iloc[METADATA_ROW_FECHA_SOLICITUD, METADATA_COL_FECHA_SOLICITUD]) else ""
                
                if len(sheet_data) > METADATA_ROW_TIENDA and len(sheet_data.columns) > METADATA_COL_TIENDA:
                    metadata.tienda = str(sheet_data.iloc[METADATA_ROW_TIENDA, METADATA_COL_TIENDA]) if pd.notna(sheet_data.iloc[METADATA_ROW_TIENDA, METADATA_COL_TIENDA]) else ""
                
                if len(sheet_data) > METADATA_ROW_ADMINISTRADOR and len(sheet_data.columns) > METADATA_COL_ADMINISTRADOR:
                    metadata.administrador = str(sheet_data.iloc[METADATA_ROW_ADMINISTRADOR, METADATA_COL_ADMINISTRADOR]) if pd.notna(sheet_data.iloc[METADATA_ROW_ADMINISTRADOR, METADATA_COL_ADMINISTRADOR]) else ""
                    
            except Exception as meta_error:
                result.errors.append(f"Error extracting metadata: {str(meta_error)}")
                self.logger.warning(f"Metadata extraction error in sheet '{sheet_name}': {str(meta_error)}")
            
            # Extract main data (columns B through I) starting from configured row
            try:
                if len(sheet_data) > DATA_START_ROW:  # Ensure we have data rows
                    # Read main data: from data start row onwards, columns B through I
                    main_data_rows = sheet_data.iloc[DATA_START_ROW:, IGNORE_COLUMN_INDEX + 1:MAIN_DATA_END_COLUMN + 1]
                    
                    # Set the first data row as headers for main data
                    if len(main_data_rows) > 0:
                        headers = main_data_rows.iloc[0]  # First row becomes headers
                        main_data_rows = main_data_rows.iloc[1:]  # Remaining rows are data
                        main_data_rows.columns = headers
                        
                        # Remove completely empty rows from main data
                        main_data_rows = main_data_rows.dropna(how='all')
                        
                        # Check for missing DNI in rows with data
                        if 'DNI' in main_data_rows.columns or any('dni' in str(col).lower() for col in main_data_rows.columns):
                            # Find DNI column (case insensitive)
                            dni_col = None
                            for col in main_data_rows.columns:
                                if 'dni' in str(col).lower():
                                    dni_col = col
                                    break
                            
                            if dni_col is not None:
                                # Check for missing DNI in non-empty rows
                                non_empty_rows = main_data_rows.dropna(how='all')
                                missing_dni_count = non_empty_rows[dni_col].isna().sum()
                                
                                if missing_dni_count > 0:
                                    result.errors.append(f"{missing_dni_count} rows with data are missing DNI")
                        else:
                            result.warnings.append("No DNI column found - cannot validate DNI completeness")
                        
                        # Extract uniform data (columns J through R) starting from row 8
                        uniform_data = None
                        if len(sheet_data) > UNIFORM_DATA_START_ROW and len(sheet_data.columns) > UNIFORM_DATA_END_COLUMN:
                            uniform_data_rows = sheet_data.iloc[UNIFORM_DATA_START_ROW:, UNIFORM_DATA_START_COLUMN:UNIFORM_DATA_END_COLUMN + 1]
                            
                            # Set column names for uniform data
                            if len(uniform_data_rows.columns) == len(UNIFORM_COLUMN_NAMES):
                                uniform_data_rows.columns = UNIFORM_COLUMN_NAMES
                                
                                # Clean uniform data the same way as main data (remove completely empty rows)
                                uniform_data_rows = uniform_data_rows.dropna(how='all')
                                
                                # Reset index to align with main data
                                uniform_data_rows = uniform_data_rows.reset_index(drop=True)
                                main_data_rows = main_data_rows.reset_index(drop=True)
                                
                                # Take only the rows that correspond to the cleaned main data
                                if len(uniform_data_rows) >= len(main_data_rows):
                                    uniform_data = uniform_data_rows.iloc[:len(main_data_rows)].copy()
                                else:
                                    uniform_data = uniform_data_rows.copy()
                                    result.warnings.append(f"Uniform data has fewer rows ({len(uniform_data_rows)}) than main data ({len(main_data_rows)})")
                            else:
                                result.warnings.append(f"Uniform data has {len(uniform_data_rows.columns)} columns, expected {len(UNIFORM_COLUMN_NAMES)}")
                        else:
                            result.warnings.append("Insufficient data for uniform columns (J-R)")
                        
                        result.data = main_data_rows
                        result.uniform_data = uniform_data
                        result.people_parsed = len(main_data_rows)  # After cleaning empty rows
                        result.total_lines = len(sheet_data)  # Total lines in sheet
                        
                        self.logger.info(f"Sheet '{sheet_name}': {result.people_parsed} people parsed from {result.total_lines} total lines")
                        if uniform_data is not None:
                            self.logger.info(f"Sheet '{sheet_name}': {len(uniform_data)} uniform data rows extracted")
                    else:
                        result.warnings.append("No data rows found after headers")
                else:
                    result.warnings.append(f"Sheet has insufficient rows (less than {DATA_START_ROW + 1})")
                    
            except Exception as data_error:
                result.errors.append(f"Error extracting data: {str(data_error)}")
                self.logger.error(f"Data extraction error in sheet '{sheet_name}': {str(data_error)}")
            
            # Validate metadata
            if not metadata.fecha_solicitud:
                result.warnings.append("Missing fecha_solicitud (C3)")
            if not metadata.tienda:
                result.warnings.append("Missing tienda (C4)")
            if not metadata.administrador:
                result.warnings.append("Missing administrador (C5)")
            
        except Exception as e:
            result.errors.append(f"Critical error parsing worksheet: {str(e)}")
            self.logger.exception(f"Critical error parsing worksheet '{sheet_name}': {str(e)}")
        
        return result
    
    def validate_excel_data(self, excel_data: ExcelData) -> ExcelValidationResult:
        """
        Validate Excel data for processing.
        
        Args:
            excel_data: ExcelData object to validate
            
        Returns:
            ExcelValidationResult with validation details
        """
        result = ExcelValidationResult(is_valid=False)
        
        try:
            if not excel_data.is_loaded:
                result.errors.append("No Excel worksheets could be loaded")
                result.message = "Excel file could not be loaded or has no valid worksheets"
                return result
            
            if excel_data.total_worksheets == 0:
                result.errors.append("Excel file contains no worksheets")
                result.message = "Excel file is empty or corrupted"
                return result
            
            # Validate each worksheet
            total_errors = 0
            total_warnings = 0
            
            for worksheet in excel_data.worksheets:
                if worksheet.errors:
                    result.errors.extend([f"Sheet '{worksheet.metadata.sheet_name}': {error}" for error in worksheet.errors])
                    total_errors += len(worksheet.errors)
                
                if worksheet.warnings:
                    result.warnings.extend([f"Sheet '{worksheet.metadata.sheet_name}': {warning}" for warning in worksheet.warnings])
                    total_warnings += len(worksheet.warnings)
            
            # Check if we have at least one successful worksheet
            if excel_data.successful_worksheets == 0:
                result.errors.append("No worksheets could be parsed successfully")
                result.message = f"All {excel_data.total_worksheets} worksheets failed to parse"
                return result
            
            # Check for occupation mapping issues
            occupation_mapping_issues = set()
            for worksheet in excel_data.worksheets:
                if worksheet.data is not None and 'cargo' in worksheet.data.columns:
                    for cargo in worksheet.data['cargo'].dropna().unique():
                        if cargo and str(cargo).strip():
                            normalized = self.unified_service.normalize_occupation(str(cargo))
                            if normalized != str(cargo).upper():
                                occupation_mapping_issues.add(f"'{cargo}' → '{normalized}'")
            
            # Add occupation mapping warnings
            if occupation_mapping_issues:
                result.warnings.append(f"Occupation mapping issues found: {', '.join(sorted(occupation_mapping_issues))}")
                result.warnings.append("Consider adding these occupations to the synonyms list in Configuration tab")
                total_warnings += len(occupation_mapping_issues) + 1
            
            # Validation passes if we have at least one successful worksheet
            result.is_valid = True
            
            # Create summary message
            success_msg = f"Excel file validated: {excel_data.successful_worksheets}/{excel_data.total_worksheets} worksheets parsed successfully"
            success_msg += f", {excel_data.total_people_parsed} people total"
            
            if total_errors > 0:
                success_msg += f", {total_errors} errors"
            if total_warnings > 0:
                success_msg += f", {total_warnings} warnings"
                
            result.message = success_msg
            
            return result
            
        except Exception as e:
            result.errors.append(f"Validation error: {str(e)}")
            result.message = "Unexpected error during validation"
            self.logger.exception(f"Excel validation error: {str(e)}")
            return result


class FileGenerationService:
    """Service for generating files from Excel data using templates."""
    
    def __init__(self, logger: logging.Logger, unified_service: UnifiedConfigService):
        self.logger = logger
        self.unified_service = unified_service
    
    
    def generate_files(self, excel_data: ExcelData, config: AppConfig, options: Optional[GenerationOptions] = None) -> GenerationResult:
        """
        Generate files from Excel data using templates.
        
        Args:
            excel_data: Loaded Excel data
            config: Application configuration
            options: Generation options
            
        Returns:
            GenerationResult with success status and details
        """
        result = GenerationResult(success=False)
        
        try:
            # Validate inputs and setup
            validation_result = self._validate_generation_inputs(excel_data, config, options)
            if not validation_result.success:
                result.message = validation_result.message
                result.errors.extend(validation_result.errors)
                return result
            
            # Set default options if not provided
            if options is None:
                options = self._create_default_options(excel_data)
            
            # Group data by locale
            tienda_to_rows = self._group_data_by_locale(excel_data, options)
            if not tienda_to_rows:
                result.message = "No matching locales or people to generate"
                return result
            
            # Generate files
            files_generated = self._generate_documents(tienda_to_rows, config, options)
            
            # Create success result
            result.success = True
            result.files_generated = files_generated
            total_people = sum(len(people) for people in tienda_to_rows.values())
            templates = self._get_enabled_templates(options)
            result.message = f"Generated {', '.join(templates)} documents for {total_people} people across {len(tienda_to_rows)} locales"
            
            self.logger.info(result.message)
            return result
            
        except Exception as e:
            error_msg = f"Error during file generation: {str(e)}"
            self.logger.exception(error_msg)
            result.message = error_msg
            result.errors.append(error_msg)
            return result

    def _validate_generation_inputs(self, excel_data: ExcelData, config: AppConfig, options: Optional[GenerationOptions]) -> GenerationResult:
        """Validate inputs for file generation."""
        result = GenerationResult(success=False)
        
        if not excel_data.is_loaded:
            result.message = "No Excel data loaded"
            return result
        
        # Validate template files
        template_errors = TemplateValidator.validate_autorizacion_template(config)
        if template_errors:
            result.message = "Template validation failed"
            result.errors.extend(template_errors)
            return result
        
        # Create and validate destination directory
        dest_path = Path(config.destination_path)
        dest_path.mkdir(parents=True, exist_ok=True)
        if not dest_path.exists() or not dest_path.is_dir():
            result.message = f"Destination path is invalid: {dest_path}"
            return result
        
        # Check if any templates are enabled
        if options and not (options.autorizacion_enabled or options.cargo_enabled):
            result.message = "No templates selected for generation"
            return result
        
        # Check docxtpl availability
        try:
            from docxtpl import DocxTemplate  # noqa: F401
        except Exception as e:
            result.message = "docxtpl not available; please install python-docx-template"
            result.errors.append(str(e))
            return result
        
        result.success = True
        return result
    
    def _create_default_options(self, excel_data: ExcelData) -> GenerationOptions:
        """Create default generation options."""
        locales = [w.metadata.tienda for w in excel_data.worksheets if w.metadata.tienda]
        return GenerationOptions(
            selected_locales=locales,
            combine_per_local=False,
            cargo_enabled=True,
            autorizacion_enabled=True
        )
    
    def _get_enabled_templates(self, options: GenerationOptions) -> List[str]:
        """Get list of enabled template names."""
        templates = []
        if options.autorizacion_enabled:
            templates.append("AUTORIZACION")
        if options.cargo_enabled:
            templates.append("CARGO")
        return templates
    
    def _group_data_by_locale(self, excel_data: ExcelData, options: GenerationOptions) -> Dict[str, List[Dict[str, Any]]]:
        """Group worksheet data by locale (tienda) with context building."""
        tienda_to_rows: Dict[str, List[Dict[str, Any]]] = {}
        
        for ws in excel_data.worksheets:
            tienda = str(ws.metadata.tienda or "").strip()
            if not tienda or tienda not in options.selected_locales:
                continue
            
            if ws.data is None or ws.data.empty:
                continue
            
            # Skip sheets with critical errors (missing fecha_solicitud, tienda, or DNI errors)
            # Skip only if tienda is missing (critical for folder structure)
            if not tienda or tienda.strip() == "":
                self.logger.warning(f"Sheet '{ws.metadata.sheet_name}' skipped: Missing tienda")
                continue
            
            # Process each person in the worksheet
            for idx, row in ws.data.iterrows():
                person_contexts = self._build_person_contexts(row, ws, options)
                if person_contexts:
                    tienda_to_rows.setdefault(tienda, []).append(person_contexts)
        
        return tienda_to_rows
    
    
    
    def _build_person_contexts(self, row: pd.Series, ws, options: GenerationOptions) -> Dict[str, Any]:
        """Build context dictionaries for a person based on enabled templates."""
        person_contexts = {}
        
        # Get uniform row for both contexts
        uniform_row = self._get_uniform_row_for_person(row, ws)
        
        if options.autorizacion_enabled:
            autorizacion_ctx = self._build_autorizacion_context(row, ws.metadata, uniform_row)
            if autorizacion_ctx:
                person_contexts["AUTORIZACION"] = autorizacion_ctx
                self.logger.debug(f"Built AUTORIZACION context for {autorizacion_ctx.get('nombre', 'unknown')}")
            else:
                self.logger.warning("Failed to build AUTORIZACION context for row")
        
        if options.cargo_enabled:
            cargo_ctx = self._build_cargo_context(row, ws.metadata, uniform_row)
            if cargo_ctx:
                person_contexts["CARGO"] = cargo_ctx
                self.logger.debug(f"Built CARGO context for {cargo_ctx.get('nombre', 'unknown')} with {len(cargo_ctx.get('prendas', []))} prendas")
            else:
                self.logger.warning("Failed to build CARGO context for row")
        
        return person_contexts
    
    def _generate_documents(self, tienda_to_rows: Dict[str, List[Dict[str, Any]]], config: AppConfig, options: GenerationOptions) -> int:
        """Generate all documents for all locales and people."""
        dest_path = Path(config.destination_path)
        files_generated = 0
        
        for tienda, people in tienda_to_rows.items():
            tienda_folder = dest_path / self._sanitize_name(tienda)
            tienda_folder.mkdir(parents=True, exist_ok=True)
            
            template_docs: Dict[str, List[Path]] = {}
            
            # Generate individual documents for each person
            for person_contexts in people:
                person_name = self._extract_person_name(person_contexts)
                if not person_name:
                    self.logger.warning("Skipping person with no name")
                    continue
                
                person_folder = tienda_folder / self._sanitize_name(person_name)
                person_folder.mkdir(parents=True, exist_ok=True)
                
                # Generate documents for each template type
                for template_type, context in person_contexts.items():
                    docx_path = self._generate_single_document(
                        template_type, context, person_folder, config
                    )
                    if docx_path:
                        template_docs.setdefault(template_type, []).append(docx_path)
                        files_generated += 1
            
            # Create combined documents if requested
            if options.combine_per_local:
                files_generated += self._create_combined_documents(template_docs, tienda_folder, tienda)
        
        return files_generated
    
    def _extract_person_name(self, person_contexts: Dict[str, Any]) -> str:
        """Extract person name from any available context."""
        for context in person_contexts.values():
            if context.get("nombre"):
                return context["nombre"]
        return ""
    
    def _generate_single_document(self, template_type: str, context: Dict[str, Any], person_folder: Path, config: AppConfig) -> Optional[Path]:
        """Generate a single document for a person."""
        try:
            if template_type == "AUTORIZACION":
                docx_path = person_folder / f"AUTORIZACION_{self._file_stub(context)}.docx"
                self.logger.info(f"Generating AUTORIZACION document: {docx_path}")
                self._render_autorizacion_doc(config.autorizacion_template_path, context, docx_path)
                return docx_path
            
            elif template_type == "CARGO":
                docx_path = person_folder / f"CARGO_{self._file_stub(context)}.docx"
                self.logger.info(f"Generating CARGO document: {docx_path}")
                self._render_cargo_doc(config.cargo_template_path, context, docx_path)
                return docx_path
            
        except Exception as e:
            self.logger.error(f"Failed to generate {template_type} document: {e}")
        
        return None
    
    def _create_combined_documents(self, template_docs: Dict[str, List[Path]], tienda_folder: Path, tienda: str) -> int:
        """Create combined documents for each template type."""
        files_generated = 0
        
        for template_type, docs in template_docs.items():
            if docs:
                combined_path = tienda_folder / f"{template_type}_COMBINED_{self._sanitize_name(tienda)}.docx"
                self._create_combined_docx(docs, combined_path)
                files_generated += 1
        
        return files_generated

    def _extract_common_person_data(self, row: pd.Series) -> Dict[str, str]:
        """Extract common person data from a row."""
        return {
            "cargo": self._find_in_row(row, ["cargo"]) or "",
            "nombre": self._find_in_row(row, ["nombre"]) or self._extract_name(row),
            "identificacion": self._find_in_row(row, ["dni"]) or ""
        }
    
    def _get_uniform_row_for_person(self, row: pd.Series, ws) -> Optional[pd.Series]:
        """Get uniform row for a person, handling edge cases safely."""
        if not self._has_valid_uniform_data(ws):
            return None
        
        row_index = row.name if hasattr(row, 'name') else 0
        if row_index < len(ws.uniform_data):
            return ws.uniform_data.iloc[row_index]
        return None
    
    def _has_valid_uniform_data(self, ws) -> bool:
        """Check if worksheet has valid uniform data."""
        return (ws.uniform_data is not None and 
                hasattr(ws, 'data') and 
                len(ws.data) > 0)
    
    def _build_autorizacion_context(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> Optional[Dict[str, Any]]:
        """Build docxtpl context for AUTORIZACION: dia, mes, anho, local, cargo, nombre, identificacion, monto(0)."""
        try:
            # Parse fecha_solicitud using the same flexible parsing as CARGO
            dia, _, anho, fecha_formatted = self._get_system_date()
            
            # Extract numeric month for AUTORIZACION (it expects MM format)
            from datetime import datetime
            now = datetime.now()
            mes = f"{now.month:02d}"
            
            # Extract common person data
            person_data = self._extract_common_person_data(row)
            
            # Format fecha as dd / MM / yyyy for AUTORIZACION template
            fecha_template = ""
            if dia and mes and anho:
                fecha_template = f"{dia} / {mes} / {anho}"
            
            # Get monto (calculated from pricing service)
            monto_value = self._get_monto_for_person(row, metadata, uniform_row)
            monto_formatted = f"S/ {monto_value:.2f}"
            
            # Debug logging for missing data
            if not person_data["identificacion"]:
                self.logger.warning(f"No DNI found for person: {person_data['nombre']}. Available columns: {list(row.index)}")
            if not person_data["nombre"]:
                self.logger.warning(f"No nombre found. Available columns: {list(row.index)}")
            if not person_data["cargo"]:
                self.logger.warning(f"No cargo found. Available columns: {list(row.index)}")
            if not fecha_template:
                self.logger.warning(f"No date found in metadata: {metadata.fecha_solicitud}")
            
            context = {
                "dia": dia,
                "mes": mes,
                "anho": anho,
                "fecha": fecha_template,  # Add formatted date for template
                "local": metadata.tienda or "",
                "cargo": str(person_data["cargo"]),
                "nombre": str(person_data["nombre"]),
                "identificacion": str(person_data["identificacion"]),
                "monto": monto_formatted,  # Formatted with Sol currency
            }
            return context
        except Exception as e:
            self.logger.error(f"Error building context: {e}")
            return None

    def _build_cargo_context(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> Optional[Dict[str, Any]]:
        """Build docxtpl context for CARGO documents with Spanish months and prenda handling."""
        try:
            # Parse fecha_solicitud with flexible date handling
            dia, mes_string, anho, fecha_string = self._get_system_date()
            
            # Extract common person data
            person_data = self._extract_common_person_data(row)
            
            # Extract talla prenda superior (second to last row data)
            talla_superior = self._extract_talla_superior(row)
            
            # Build prendas list from uniform data
            # Use uniform_row if available, otherwise fall back to main row
            data_row = uniform_row if uniform_row is not None else row
            prendas = self._build_prendas_list(data_row, talla_superior)
            
            # Get monto (calculated from pricing service)
            monto_value = self._get_monto_for_person(row, metadata, uniform_row)
            monto_formatted = f"S/ {monto_value:.2f}"
            
            context = {
                "dia": dia,
                "mes_string": mes_string,
                "anho": anho,
                "fecha": fecha_string,
                "nombre": str(person_data["nombre"]),
                "prendas": prendas,
                "monto": monto_formatted  # Add pricing information
            }
            
            return context
        except Exception as e:
            self.logger.error(f"Error building CARGO context: {e}")
            return None

    def _render_autorizacion_doc(self, template_path: str, context: Dict[str, Any], output_docx: Path) -> None:
        """Render AUTORIZACION document."""
        self._render_document(template_path, context, output_docx)
    
    def _render_cargo_doc(self, template_path: str, context: Dict[str, Any], output_docx: Path) -> None:
        """Render CARGO document."""
        self._render_document(template_path, context, output_docx)
    
    def _render_document(self, template_path: str, context: Dict[str, Any], output_docx: Path) -> None:
        """Generic document rendering method."""
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save(str(output_docx))

    def _create_combined_docx(self, individual_docs: List[Path], output_path: Path) -> None:
        """Combine multiple DOCX files into one concatenated document using docxcompose."""
        try:
            from docxcompose.composer import Composer
            from docx import Document
            
            if not individual_docs:
                self.logger.warning("No individual documents to combine")
                return
            
            # Use the first document as master
            first_doc_path = individual_docs[0]
            if not first_doc_path.exists():
                self.logger.error(f"First document does not exist: {first_doc_path}")
                return
            
            # Load master document
            master_doc = Document(str(first_doc_path))
            composer = Composer(master_doc)
            
            # Append remaining documents
            for doc_path in individual_docs[1:]:
                if not doc_path.exists():
                    self.logger.warning(f"Document does not exist, skipping: {doc_path}")
                    continue
                    
                try:
                    # Load document to append
                    doc_to_append = Document(str(doc_path))
                    composer.append(doc_to_append)
                    self.logger.debug(f"Successfully appended: {doc_path}")
                    
                except Exception as e:
                    self.logger.warning(f"Failed to append document {doc_path}: {e}")
                    continue
            
            # Save combined document
            composer.save(str(output_path))
            self.logger.info(f"Created combined document using docxcompose: {output_path}")
            
        except ImportError:
            self.logger.error("docxcompose not available. Please install: pip install docxcompose")
            # Fallback to simple concatenation
            self._create_fallback_combined_docx(individual_docs, output_path)
            
        except Exception as e:
            self.logger.error(f"Failed to create combined document with docxcompose: {e}")
            # Fallback to simple concatenation
            self._create_fallback_combined_docx(individual_docs, output_path)
    
    def _create_fallback_combined_docx(self, individual_docs: List[Path], output_path: Path) -> None:
        """Fallback method for combining documents when docxcompose fails."""
        try:
            from docx import Document
            
            # Create a simple document listing all individual files
            fallback_doc = Document()
            fallback_doc.add_heading(f"Combined Document - {output_path.stem}", 0)
            fallback_doc.add_paragraph(f"This document combines {len(individual_docs)} individual authorization documents.")
            fallback_doc.add_paragraph("Individual documents:")
            
            for i, doc_path in enumerate(individual_docs, 1):
                if doc_path.exists():
                    fallback_doc.add_paragraph(f"{i}. {doc_path.name}")
                else:
                    fallback_doc.add_paragraph(f"{i}. {doc_path.name} (missing)")
            
            fallback_doc.add_paragraph("")
            fallback_doc.add_paragraph("Note: Please check individual documents for complete content.")
            fallback_doc.save(str(output_path))
            self.logger.info(f"Created fallback combined document: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to create fallback combined document: {e}")
    
    def _get_monto_for_person(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> float:
        """Get the monto (amount) for a person based on their uniform requirements."""
        try:
            # Get cargo and normalize it
            cargo = self._find_in_row(row, ["cargo"]) or ""
            if not cargo:
                self.logger.warning("No cargo found for person, using default pricing")
                cargo = "MOZO"  # Default cargo
            
            # Normalize cargo using synonyms
            normalized_cargo = self.unified_service.normalize_occupation(cargo)
            
            # Check if occupation was properly mapped
            if normalized_cargo != cargo.upper():
                self.logger.warning(f"⚠️  OCCUPATION MAPPING: '{cargo}' → '{normalized_cargo}' - consider adding '{cargo}' to synonyms list in Configuration tab")
            
            # Get local/tienda
            local = metadata.tienda or "OTHER"
            
            # Build prendas list for pricing calculation
            talla_superior = self._extract_talla_superior(row)
            prendas = self._build_prendas_list(uniform_row if uniform_row is not None else row, talla_superior)
            
            # Debug: log detailed information for PACKER and MOTORIZADO
            person_name = self._extract_name(row)
            if normalized_cargo in ['PACKER', 'MOTORIZADO'] or 'PACKER' in cargo.upper() or 'MOTORIZADO' in cargo.upper():
                self.logger.info(f"DEBUG PACKER/MOTORIZADO - {person_name}:")
                self.logger.info(f"  Original cargo: '{cargo}'")
                self.logger.info(f"  Normalized cargo: '{normalized_cargo}'")
                self.logger.info(f"  Local: '{local}'")
                self.logger.info(f"  Talla superior: '{talla_superior}'")
                self.logger.info(f"  Uniform row available: {uniform_row is not None}")
                if uniform_row is not None:
                    self.logger.info(f"  Uniform row columns: {list(uniform_row.index)}")
                    self.logger.info(f"  Uniform row values: {uniform_row.to_dict()}")
                self.logger.info(f"  Found {len(prendas)} prendas:")
                for i, prenda in enumerate(prendas):
                    self.logger.info(f"    {i+1}. {prenda}")
            
            # Calculate total price
            total_price = self.unified_service.calculate_total_price(prendas, normalized_cargo, local)
            
            self.logger.info(f"Calculated monto for {person_name}: "
                           f"Cargo={normalized_cargo}, Local={local}, Prendas={len(prendas)}, Total={total_price}")
            
            # Debug: log individual prenda prices
            for prenda in prendas:
                self.logger.debug(f"  Prenda: {prenda['string']}, Qty: {prenda['qty']}")
            
            return total_price
            
        except Exception as e:
            self.logger.error(f"Failed to calculate monto for person: {e}")
            return 0.0
    
    def _get_system_date(self) -> tuple:
        """Get system date in the required format."""
        from datetime import datetime
        now = datetime.now()
        dia = f"{now.day:02d}"
        mes_string = SPANISH_MONTHS.get(now.month, "")
        anho = f"{now.year}"
        fecha_string = f"{dia} de {mes_string} de {anho}"
        return dia, mes_string, anho, fecha_string
    
    def _extract_talla_superior(self, row: pd.Series) -> str:
        """Extract talla prenda superior from row data."""
        # Look for talla-related columns (second to last priority)
        talla_keys = ["talla prenda superior", "talla", "size", "talla_superior"]
        talla = self._find_in_row(row, talla_keys) or ""
        return str(talla).strip().upper()
    
    def _build_prendas_list(self, row: pd.Series, talla_superior: str) -> List[Dict[str, Any]]:
        """Build list of prendas from uniform data with quantities."""
        prendas = []
        
        # Debug: log available columns for PACKER/MOTORIZADO debugging
        if any(name in str(row.index).upper() for name in ['PACKER', 'MOTORIZADO']):
            self.logger.info(f"DEBUG _build_prendas_list - Available columns: {list(row.index)}")
            self.logger.info(f"DEBUG _build_prendas_list - Looking for: {UNIFORM_COLUMN_NAMES}")
        
        # Iterate through uniform columns to find quantities
        for prenda_name in UNIFORM_COLUMN_NAMES:
            # Try to find the prenda column directly first
            qty_value = None
            if prenda_name in row.index:
                qty_value = row[prenda_name]
                # Handle Series values (from duplicate column names)
                if isinstance(qty_value, pd.Series):
                    qty_value = qty_value.iloc[0] if len(qty_value) > 0 else None
            else:
                # Fallback to the general search method
                qty_value = self._find_in_row(row, [prenda_name])
            
            # Check if we have a valid quantity value (not NaN, not empty, not zero)
            if (qty_value is not None and 
                pd.notna(qty_value) and 
                str(qty_value).strip() not in ['', 'nan', 'NaN', '0']):
                try:
                    qty = int(float(str(qty_value).strip()))
                    if qty > 0:
                        # Format prenda string: keep prefixes for pricing, remove only for display
                        clean_name = prenda_name.upper()
                        display_name = clean_name
                        
                        # For display, remove prefixes
                        if clean_name.startswith('PACKER'):
                            display_name = clean_name.replace('PACKER', '').strip()
                        elif clean_name.startswith('DELIVERY'):
                            display_name = clean_name.replace('DELIVERY', '').strip()
                        
                        # Create formatted prenda string for display
                        prenda_string = f"{display_name} TALLA {talla_superior}" if display_name not in {"ANDARIN", "MANDILON", "GORRA"} else display_name
                        
                        prenda_dict = {
                            "string": prenda_string,
                            "qty": qty,
                            "prenda_type": clean_name  # Keep original name for pricing
                        }
                        prendas.append(prenda_dict)
                        
                except (ValueError, TypeError):
                    # Only log warning for truly invalid values, not NaN/empty
                    if str(qty_value).strip() not in ['', 'nan', 'NaN']:
                        self.logger.warning(f"Invalid quantity value for {prenda_name}: {qty_value}")
                    continue
        
        return prendas

    def _find_in_row(self, row: pd.Series, keys: List[str]) -> Optional[str]:
        """Find a value in a row by searching for column names that contain any of the specified keys."""
        lowered = {str(k).lower(): k for k in row.index}
        
        # First try exact matches
        for needle in keys:
            if needle.lower() in lowered:
                val = row[lowered[needle.lower()]]
                # Handle Series values (from duplicate column names)
                if isinstance(val, pd.Series):
                    val = val.iloc[0] if len(val) > 0 else None
                if pd.notna(val) and str(val).strip():
                    return str(val).strip()
        
        # Then try partial matches
        for key in lowered:
            for needle in keys:
                if needle.lower() in key:
                    val = row[lowered[key]]
                    # Handle Series values (from duplicate column names)
                    if isinstance(val, pd.Series):
                        val = val.iloc[0] if len(val) > 0 else None
                    if pd.notna(val) and str(val).strip():
                        return str(val).strip()
        
        return None

    def _extract_name(self, row: pd.Series) -> str:
        lowered = {str(k).lower(): k for k in row.index}
        combined = None
        for key in lowered:
            if "nombre" in key and "apellido" in key:
                combined = row[lowered[key]]
                # Handle Series values (from duplicate column names)
                if isinstance(combined, pd.Series):
                    combined = combined.iloc[0] if len(combined) > 0 else None
                break
        if pd.notna(combined) and str(combined).strip():
            return str(combined).strip()
        first = None
        last = None
        for key in lowered:
            if first is None and ("nombre" in key or "name" in key):
                first = row[lowered[key]]
                # Handle Series values (from duplicate column names)
                if isinstance(first, pd.Series):
                    first = first.iloc[0] if len(first) > 0 else None
            if last is None and ("apellido" in key or "last" in key):
                last = row[lowered[key]]
                # Handle Series values (from duplicate column names)
                if isinstance(last, pd.Series):
                    last = last.iloc[0] if len(last) > 0 else None
        name = ""
        if pd.notna(first):
            name = str(first).strip()
        if pd.notna(last):
            name = (name + " " + str(last).strip()).strip()
        return name

    def _sanitize_name(self, s: str) -> str:
        return "".join(ch for ch in s if ch.isalnum() or ch in ("_", "-", " ")).strip().replace(" ", "_")

    def _file_stub(self, ctx: Dict[str, Any]) -> str:
        parts = [ctx.get("nombre", "").strip(), ctx.get("cargo", "").strip()]
        return self._sanitize_name("_".join([p for p in parts if p]))


# ConfigService removed - use ConfigManager directly from config_manager module
